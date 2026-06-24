#!/usr/bin/env python3
"""Generate binary test fixtures for PDF, DOCX, and Excel.

Run once (or after changing fixture content):
  uv run --group dev tests/generate_fixtures.py

Outputs files to tests/data/{pdf,docx,excel}/.
"""

from pathlib import Path

DATA_DIR = Path(__file__).parent / "data"


def generate_pdf():
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=14)
    pdf.cell(0, 10, "Knowledge Project Skills - Test Document",
             new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 8, (
        "This is a test PDF used to verify the PDF preprocessor. "
        "It contains multiple paragraphs with entities and facts."
    ))
    pdf.ln(4)
    pdf.multi_cell(0, 8, (
        "Alice founded Acme Corp in Paris in 2010. "
        "The company grew to 500 employees by 2020 and was acquired by GlobalTech in 2023."
    ))
    pdf.add_page()
    pdf.set_font("Helvetica", "B", size=12)
    pdf.cell(0, 10, "Second Page", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 8, "Key fact: revenue reached 50M EUR in 2022.")

    out = DATA_DIR / "pdf/sample.pdf"
    out.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(out))
    print(f"  wrote {out}")


def generate_docx():
    import docx

    doc = docx.Document()
    doc.add_heading("Knowledge Project Skills — Test Document", level=1)
    doc.add_paragraph(
        "This is a test Word document used to verify the DOCX preprocessor. "
        "It contains headings, paragraphs, and a table."
    )
    doc.add_heading("Background", level=2)
    doc.add_paragraph(
        "Alice founded Acme Corp in Paris in 2010. "
        "Bob joined as CTO in 2012. The company raised Series A in 2015."
    )
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    headers = ["Name", "Role", "Year"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "Alice"
    table.rows[1].cells[1].text = "CEO"
    table.rows[1].cells[2].text = "2010"
    table.rows[2].cells[0].text = "Bob"
    table.rows[2].cells[1].text = "CTO"
    table.rows[2].cells[2].text = "2012"

    out = DATA_DIR / "docx/sample.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  wrote {out}")


def generate_excel():
    import openpyxl

    wb = openpyxl.Workbook()

    ws1 = wb.active
    ws1.title = "People"
    ws1.append(["name", "age", "city", "score"])
    ws1.append(["Alice", 30, "Paris", 95.5])
    ws1.append(["Bob", 25, "Lyon", 87.3])
    ws1.append(["Carol", 35, "Berlin", 92.1])

    ws2 = wb.create_sheet("Events")
    ws2.append(["date", "event", "location"])
    ws2.append(["2010-01-15", "Company founded", "Paris"])
    ws2.append(["2015-06-01", "Series A", "London"])
    ws2.append(["2023-03-10", "Acquisition", "New York"])

    out = DATA_DIR / "excel/sample.xlsx"
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out))
    print(f"  wrote {out}")


def generate_rulebook_pdf():
    """A mini-rulebook with a rule-procedure section and a lookup table.

    Designed to test extraction of entity `body` content:
    - Combat: multi-paragraph procedure with sub-rules and a worked example.
      Should become a concept entity with a `body`.
    - Wounds: a named d6 lookup table.
      Should become a concept entity whose `body` contains the full table.
    - Shallow entities (Aldric, Ironforge, The Depths) need only `context`.
    """
    from fpdf import FPDF, FontFace

    pdf = FPDF()

    # --- Page 1: title and overview ---
    pdf.add_page()
    pdf.set_font("Helvetica", "B", size=16)
    pdf.cell(0, 12, "Ironforge Quickstart Rules", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 7, (
        "Ironforge is a rules-lite adventure game for 2 to 5 players, "
        "designed by Aldric. One player is the Guide; the rest play characters "
        "exploring The Depths, a vast underground ruin."
    ))
    pdf.ln(4)
    pdf.multi_cell(0, 7, (
        "Characters have three ability scores: Strength (STR), Dexterity (DEX), "
        "and Willpower (WIL), each rolled 3d6 at creation. Starting Hit Points "
        "(HP) are rolled 1d6. A character with STR 0 dies; DEX 0 means paralysis; "
        "WIL 0 means madness."
    ))
    pdf.ln(4)
    pdf.multi_cell(0, 7, (
        "Ironforge was first published in 2021. It is released under CC-BY 4.0."
    ))

    # --- Page 2: Combat ---
    pdf.add_page()
    pdf.set_font("Helvetica", "B", size=14)
    pdf.cell(0, 10, "Combat", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 7, (
        "Combat has no fixed initiative. Each round, all participants declare "
        "their intentions and the Guide resolves them in the order that makes "
        "fictional sense."
    ))
    pdf.ln(3)
    pdf.set_font("Helvetica", "B", size=11)
    pdf.cell(0, 7, "Attacking", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 7, (
        "Roll your weapon's damage die and subtract the target's Armor value. "
        "The remainder is dealt as damage to HP. Unarmed strikes always deal 1d4."
    ))
    pdf.ln(3)
    pdf.set_font("Helvetica", "B", size=11)
    pdf.cell(0, 7, "Multiple Attackers", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 7, (
        "When two or more attackers target the same enemy in the same round, "
        "each rolls their damage die but only the single highest result is applied. "
        "Extra attackers improve the odds of a high roll but do not stack damage."
    ))
    pdf.ln(3)
    pdf.set_font("Helvetica", "B", size=11)
    pdf.cell(0, 7, "Impaired and Enhanced Attacks", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 7, (
        "An attack is Impaired when the attacker is at a clear disadvantage "
        "(prone, blinded, using a weapon in an awkward space). Roll d4 regardless "
        "of weapon. An attack is Enhanced when the attacker has a decisive "
        "advantage. Roll d12 regardless of weapon."
    ))
    pdf.ln(3)
    pdf.set_font("Helvetica", "B", size=11)
    pdf.cell(0, 7, "Morale", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 7, (
        "Enemy groups must pass a WIL save to avoid fleeing when they suffer "
        "their first casualty, and again when they lose half their number. "
        "Lone enemies check morale when reduced below half HP."
    ))
    pdf.ln(3)
    pdf.set_font("Helvetica", "BI", size=11)
    pdf.cell(0, 7, "Example", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "I", size=11)
    pdf.multi_cell(0, 7, (
        "Three players attack a Stone Golem (2 Armor). They roll d6, d8, and d6 "
        "and get 3, 7, and 5. Only the 7 is applied: 7 minus 2 Armor = 5 damage "
        "to the Golem. The other two rolls are discarded."
    ))

    # --- Page 3: Wounds table ---
    pdf.add_page()
    pdf.set_font("Helvetica", "B", size=14)
    pdf.cell(0, 10, "Wounds", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 7, (
        "When an attack reduces a character's HP to exactly 0, roll 1d6 and "
        "consult the Wounds table. The result number equals the damage dealt by "
        "that attack (minimum 1, maximum 6)."
    ))
    pdf.ln(3)

    table_data = [
        ("d6", "Wound"),
        ("1", "Bruised: lose 1 HP permanently until you receive proper rest."),
        ("2", "Winded: you cannot take an action on your next turn."),
        ("3", "Bleeding: lose 1 HP at the start of each round until tended."),
        ("4", "Concussed: all your rolls are Impaired until you rest safely."),
        ("5", "Broken Limb: one of your hands is unusable until magically healed."),
        ("6", "Mortal Wound: you fall unconscious and die in 1 hour without aid."),
    ]

    pdf.set_font("Helvetica", size=10)
    header_style = FontFace(emphasis="BOLD")
    with pdf.table(col_widths=(20, 150), borders_layout="ALL", line_height=6) as table:
        for i, (roll, effect) in enumerate(table_data):
            row = table.row()
            style = header_style if i == 0 else None
            row.cell(roll, style=style)
            row.cell(effect, style=style)

    pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 7, (
        "A character reduced to STR 0 by wounds dies outright, bypassing this table."
    ))

    out = DATA_DIR / "pdf/rules-quickstart.pdf"
    out.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(out))
    print(f"  wrote {out}")


if __name__ == "__main__":
    print("Generating binary test fixtures...")
    generate_pdf()
    generate_docx()
    generate_excel()
    generate_rulebook_pdf()
    print("Done.")
